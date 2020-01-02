import os.path
import sys
import re
import os
import difflib
import docx

RED = '\033[0;31m'
GREEN = '\033[0;32m'
YELLOW = '\033[0;33m'


class NewProjectDocument:
    def __init__(self, currentPath):
        self._currentPath = currentPath
        self._pathSearch = False
        self._editDocument = False
        self._addParagraph = False
        self._currentParagraph = ""
        self._currentSentence = ""
        self._makeSentence = False
    # end of init

    def returnPath(self):
        return(self._currentPath)
    # end of returnPath

    def getAdd(self):
        return(self._addParagraph)
    #end of returnAdd

    def setEditDocument(self, userInput):
        self._editDocument = userInput
    # end of setEditDocument

    def setPathSearch(self, uInput):
        self._pathSearch = uInput
    # end of setPathSearch

    def clearPath(self):
        self._currentPath = defaultPath
    # end of clearPath

    def getPathSearch(self):
        return(self._pathSearch)
    # end of getPathSearch

    def editFile(self):
        return(self._pathSearch)
    # end of

    def getEdit(self):
        return(self._editDocument)
    # end of getEdit

    def addPath(self, addedPath):
        if re.search(r'\b(open file)\b', addedPath, re.I):
            os.startfile(self._currentPath)
            self.setPathSearch(False)
        elif addedPath == "stop search":
            self.setPathSearch(False)
            sys.stdout.write(RED)
            print('File Search Has Stopped')
        else:
            possiblePaths = os.listdir(self._currentPath)
            try:
                test = (difflib.get_close_matches(
                    addedPath, possiblePaths))
                result = test[0]

                if os.path.isdir(self._currentPath + "\\" + result):
                    sys.stdout.write(GREEN)
                    print("Path Detected")
                    self._currentPath = self._currentPath + "\\" + result
                    print("Current Path is:", self.returnPath())

                elif os.path.isfile(self._currentPath + "\\" + result):
                    sys.stdout.write(GREEN)
                    print("Document Found")
                    print("Say Open File to View in Word")
                    print("Say Edit File to View and Edit in Program")
                    self._currentPath = self._currentPath + "\\" + result
                    print("Current Path is:", self.returnPath())
                    self._pathSearch = False
                # end of if
            except:
                sys.stdout.write(RED)
                print("Path Not Detected")
            # end of try and except
        # end of if
    # end of addPath

    def setUp(self):
        self._editDocument = True
        self._document = docx.Document(self._currentPath)
        sys.stdout.write(GREEN)
        print("File Set Up Completed")
    # end of setUp

    def setAdd(self, userInput):
        self._addParagraph = userInput
    #end of setAdd

    def showText(self):
        document = self._document
        fullText = []
        for para in document.paragraphs:
            fullText.append(para.text)
        # end of for
        for text in fullText:
            sys.stdout.write(YELLOW)
            print(text)
        # end of for
    #end of showText

    def saveDocument(self):
        self._document.save(self._currentPath)
    #end of saveDocument

    def addParagraph(self, currentData):
        if re.search(r'\b(New Sentence Create)\b', currentData, re.I):
            self._makeSentence = True
            sys.stdout.write(GREEN)
            print("Currently in Append Mode")
            print("New Sentence Has Been Created")
        elif re.search(r'\b(Sentence Finish)\b', currentData, re.I) and self._makeSentence:
            self._currentSentence += "."
            self._currentParagraph += self._currentSentence
            self._currentSentence = ""
            self._makeSentence = False
            print(self._makeSentence)
            print("Sentence Has Been Added")
            print("Current Paragraph:", self._currentParagraph)
        elif re.search(r'\b(Delete Previous)\b', currentData, re.I) and self._makeSentence:
            tempList = self._currentSentence.split()
            tempList = tempList[:-1]
            self._currentSentence = ' '.join(word for word in tempList)
            sys.stdout.write(GREEN)
            print("Previous Word Deleted")
            print("Current Sentence:", self._currentSentence)
        elif self._makeSentence:
            self._currentSentence += " "
            self._currentSentence += currentData
            sys.stdout.write(GREEN)
            print("Current Sentence:", self._currentSentence)
        elif not(self._makeSentence) and re.search(r'\b(Finish Paragraph)\b', currentData, re.I):
            self._inputParagraph = self._document.add_paragraph(self._currentParagraph)
            sys.stdout.write(GREEN)
            print("Paragraph Inputted")
            print("Current Paragraph Is:", self._currentParagraph)
            self._currentParagraph = ""
        elif re.search(r'\b(Stop Edit)\b', currentData, re.I) and not(self._makeSentence):
            self._currentParagraph = ""
            self._inputParagraph = None
            self.setAdd(False)
            sys.stdout.write(GREEN)
            print("Paragraph Edit Stopped")
        else:
            sys.stdout.write(RED)
            print("Command Not Accepted")
        #end of if
    #end of addParagraph 

        
# end of NewProjectDocument
